import sys
from PyQt5.QtWidgets import QApplication,  QMainWindow,  QMessageBox, QFileDialog, QTableWidgetItem
from qt_ui.ControlCenter import  *
from YUSCO.Core.DB_RDB import RDBConn
from YUSCO.Core.DB_ORACLE import OracleDB_dic
from YUSCO.Core.DB_SQL import SQLConn
from YUSCO.Util.mail_sender import SendMail
from pyodbc import connect as pyodbc_connect
import cx_Oracle

from subprocess import call as subprocess_call
import socket
import csv

import xlwt
from xlwt import Workbook, easyxf
import datetime

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm,Pt,RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH


class TinyControlCenter(QMainWindow,  Ui_TinyControlCenter):
    def __init__(self,  parent=None):
        super(TinyControlCenter,  self).__init__(parent)
        self.setupUi(self)
        self.LANButton.clicked.connect(self.LANConnect)
        self.WLANButton.clicked.connect(self.WLANConnect)
        self.btn_LoadData.clicked.connect(self.Load_Data)
        self.btn_LineBot.clicked.connect(self.LineBotAction)
        self.btn_SaveCSV.clicked.connect(self.SaveCSV)
        self.btn_SaveXLS.clicked.connect(self.SaveXLS)
        self.btn_SendMail.clicked.connect(self.MailSend)
        self.btn_RepData.clicked.connect(self.ReportData)
        self.btn_RepExcel.clicked.connect(self.savefile)

        self.dateEdit.setDate(QtCore.QDate.currentDate().addDays(-7))
        self.dateEdit_2.setDate(QtCore.QDate.currentDate())

    def LANConnect(self):
        subprocess_call('D:\TinyCMD\origin_setting.bat')

    def WLANConnect(self):
        subprocess_call('D:\TinyCMD\wifi_setting.bat')

    def Load_Data(self):

        if check_ip():
            s_sql = self.lineEdit.text().strip()
            db_name = str(self.comboBox.currentText())
            if s_sql :
                conn = pyodbc_connect(RDBConn(db_name))
#                s_sql = "select order_no_item, order_thick_max, order_thick_min from ordb011m limit to 10 rows "
                result = list(conn.execute(s_sql))
                
#                print(result)
#                print(len(result[0]))
                self.tableWidget.setRowCount(0)
                self.tableWidget.setColumnCount(len(result[0]))
                for row_number ,  row_data in enumerate(result):
                    self.tableWidget.insertRow(row_number)
                    for column_number ,  data in enumerate(row_data):
                        self.tableWidget.setItem(row_number,  column_number,  QtWidgets.QTableWidgetItem(str(data)))
                conn.close()
            else:
                QMessageBox.about(self, "Warning:", "請勿輸入空白")
        else:
            QMessageBox.about(self,"Warning:", "請先切換回內網")

    def LineBotAction(self):

        if check_ip():

            LineBotChoice = str(self.comboBox_2.currentText())
    #        QMessageBox.about(self,"Info:", LineBotChoice)

            conn = cx_Oracle.connect(OracleDB_dic('RP547A_TQC'))
            s_sql = "select empl_no,line_userid from linebot_user order by empl_no "

            cursor = conn.cursor()
            cursor.execute(s_sql)
            list_emplno =[]
            for result in cursor.fetchall():
                list_emplno.append(list(result))
            cursor.close()
            conn.close()

            try:
                conn = pyodbc_connect(RDBConn('PUB'))
                for i1, inner_l in enumerate(list_emplno):
                    emplno = str(list_emplno[i1][0].strip())
                    s_sql = "select name from pamf01 where emplno ='" + emplno + "'"
                    result = list(conn.execute(s_sql))
                    s_name = result[0][0]
                    list_emplno[i1].append(s_name)

                conn.close()
            except Exception as e:
                print('Error: something worng, except message : ' + str(e))

    #        print(list_emplno)

            self.tableWidget_2.setRowCount(0)
            self.tableWidget_2.setColumnCount(len(list_emplno[0]))

            for row_number ,  row_data in enumerate(list_emplno):
                self.tableWidget_2.insertRow(row_number)
                for column_number ,  data in enumerate(row_data):
                    self.tableWidget_2.setItem(row_number,  column_number,  QtWidgets.QTableWidgetItem(str(data)))
        else:
            QMessageBox.about(self,"Warning:", "請先切換回內網")

    def SaveXLS(self):
        path = QFileDialog.getSaveFileName(self, 'Save File', '', 'XLS(*.xls)')
        if path[0] != '':
            fileName = path[0]
            self.statusbar.showMessage(fileName)

            # Workbook is created
            wb = Workbook()
            
            # add_sheet is used to create sheet.
            sheet1 = wb.add_sheet('Sheet 1')
            for rowNumber in range(self.tableWidget.rowCount()):
                for columnNumber in range(self.tableWidget.columnCount()):
                    item = self.tableWidget.item(rowNumber,columnNumber)
                    sheet1.write(rowNumber, columnNumber, item.text())
            wb.save(fileName)

    def SaveCSV(self):
        path = QFileDialog.getSaveFileName(self, 'Save File', '', 'CSV(*.csv)')
        if path[0] != '':
            fileName = path[0]
            self.statusbar.showMessage(fileName)

            with open(fileName, 'w') as stream:
                writer = csv.writer(stream)
                for rowNumber in range(self.tableWidget.rowCount()):
                    rowdata =[]
                    for columnNumber in range(self.tableWidget.columnCount()):
                        item = self.tableWidget.item(rowNumber,columnNumber)
                        if item :
                            rowdata.append(item.text())
                        else:
                            rowdata.append('')
                    writer.writerow(rowdata)

    def MailSend(self):
        QMessageBox.about(self,"Info:", "Send Mail")
        count = 0
        while (count<100):
            curr_dt = datetime.datetime.now().strftime("%Y/%m/%d %H:%M:%S")
            SendMail(curr_dt, count)
            count = count + 1



    def ReportData(self):
#        QMessageBox.about(self,"Info:", "Report Data")
        conn = pyodbc_connect(SQLConn('NTSR12','PER'))

#        date_s = self.dateEdit.date().toPyDate()
        date_s = self.dateEdit.date().toString("yyyyMMdd")
        date_e = self.dateEdit_2.date().toString("yyyyMMdd")
        print(date_s)
        print(date_e)
#        s_sql = "SELECT emplno,sys_no,title,content,event_date,fill_date,fill_time FROM yd2a600m where fill_date between '20200401' and '20200408'"
#        s_sql = "SELECT top 2 emplno,sys_no FROM yd2a600m where fill_date between '20200327' and '20200401' "
        s_sql = " select * from (select" \
            " ( select username from secuser where userno = 'yu' + emplno) as username, " \
            " title, content, event_date, emplno, sys_no, useful," \
            " ( select deptid from yd2a110m where emplno = yd2a600m.emplno) as deptid " \
            " from yd2a600m " \
            " where fill_date between '" + date_s + "' and '" + date_e + "' ) as a " \
            " where a.deptid like 'F231%' " \
            " order by emplno, event_date "
        print(s_sql)

        result = list(conn.execute(s_sql))        
        conn.close()

        self.tableWidget_3.setRowCount(0)
        self.tableWidget_3.setColumnCount(len(result[0]) + 1)
        self.tableWidget_3.setHorizontalHeaderLabels(['選擇','姓名','標題','內容','事件日期','工號','系統','效益','部門']) 
        for row_number ,  row_data in enumerate(result):
            self.tableWidget_3.insertRow(row_number)
            chkBoxItem = QtWidgets.QCheckBox()
            self.tableWidget_3.setCellWidget(row_number , 0 ,chkBoxItem)    
            for column_number ,  data in enumerate(row_data):
                self.tableWidget_3.setItem(row_number,  column_number + 1,  QtWidgets.QTableWidgetItem(str(data)))
        self.tableWidget_3.resizeColumnsToContents()


    def savefile(self):

        date_s = self.dateEdit.date().toString("yyyyMMdd")
        date_e = self.dateEdit_2.date().toString("yyyyMMdd")
        s_title = "F231工作狀況" + date_s + " TO " + date_e

        filename,_ = QFileDialog.getSaveFileName(self, 'Save File', s_title , ".xls(*.xls)")
        wbk = xlwt.Workbook()
        sheet = wbk.add_sheet("sheet", cell_overwrite_ok=True)
        style = xlwt.XFStyle()
        font = xlwt.Font()
        font.bold = True
        style.font = font
        model = self.tableWidget_3.model()

        sheet.col(0).width = 3000
        sheet.col(1).width = 13500
        sheet.col(2).width = 2000
        sheet.col(3).width = 5000

        sheet.write(0,0, s_title )

        sheet.write(1,0, "姓名", set_style("white", "dark_green_ega"))
        sheet.write(1,1, "工作項目", set_style("white", "dark_green_ega"))
        sheet.write(1,2, "進度", set_style("white", "dark_green_ega"))
        sheet.write(1,3, "備註", set_style("white", "dark_green_ega"))

        for r in range(model.rowCount()):
            sheet.write(r+2, 0, model.data(model.index(r,1)), set_style("blue", "gold"))
            sheet.write(r+2, 1, model.data(model.index(r,2)), set_style("blue", "gold"))
            sheet.write(r+2, 2, "OK", set_style("blue", "gold"))
            sheet.write(r+2, 3, model.data(model.index(r,7)), set_style("blue", "gold"))

        sheet.write(r+4, 0, "報告項目" )
        ct = r+4
        num = 1 
        listA = []
        for r in range(model.rowCount()):
            ckBox = self.tableWidget_3.cellWidget(r,0)
            if ckBox.isChecked():
                sheet.write(ct, 1, str(num) + "." + model.data(model.index(r,2)))
                listA.append(str(num) + "." + model.data(model.index(r,2)))
                ct = ct + 1
                num = num + 1 
        wbk.save(filename)


        d=Document()


        date_t = (datetime.datetime.now() + datetime.timedelta(days=1)).strftime("%Y/%m/%d")
#        date_t = datetime.date.today().strftime("%Y/%m/%d") 
        p = d.add_paragraph('')
        p.add_run(date_t + ' 資訊部工作報告').bold = True
        p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER

        table = d.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
#        hdr_cells[0].text = "課別"
#        hdr_cells[1].text = "工作項目內容描述"
        hdr_cells[0].paragraphs[0].add_run('課別').bold=True
        hdr_cells[1].paragraphs[0].add_run('工作項目內容描述').bold=True
        table.cell(0,0).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0,1).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER

        for content in listA:
            row_cells = table.add_row().cells
            row_cells[0].text = "生產資訊一課"
            row_cells[1].text = content
        

        for cell in table.columns[0].cells:
            cell.width = Inches(1.5)
        for cell in table.columns[1].cells:
            cell.width = Inches(5)

        d.add_page_break()

        d.save('資訊部工作報告.docx')


#        for c in range(model.columnCount()):
#            text = model.headerData(c, QtCore.Qt.Horizontal)
#            sheet.write(0, c+1, text, style=style)

#        for r in range(model.rowCount()):
#            text = model.headerData(r, QtCore.Qt.Vertical)
#            sheet.write(r+1, 0, text, style=style)

#        ct = 1
#        for r in range(model.rowCount()):
#            ckBox = self.tableWidget_3.cellWidget(r,0)
#            if ckBox.isChecked():
#                for c in range(model.columnCount()):
#                    text = model.data(model.index(r, c))
#                    sheet.write(ct, c+1, text)
#                ct = ct + 1 
#        wbk.save(filename)


# item cell 的差異
#        for i in range(self.tableWidget_3.rowCount()):
#            thing = self.tableWidget_3.item(i,1)
#            thing2 = self.tableWidget_3.cellWidget(i,0)
#            if thing is not None and thing.text() != '':
#                print(thing.text())            
#            if thing2.isChecked():
#                print("aa")

def set_style(fontColor, colour):
    return easyxf(
        'font: bold 1, color %s; pattern: pattern solid,  fore_colour %s;' % (fontColor, colour) 
    )


def check_ip():
    myname = socket.getfqdn(socket.gethostname())
    myaddr = socket.gethostbyname(myname)
#    雙網後 IP 取得有問題,尚無法解,先行取消IP check
#    if myaddr[:3] == '172':
#        return True
#    else:
#        return False
    return True
    
if __name__ == "__main__":
    app = QApplication(sys.argv)
    myWin = TinyControlCenter()
    myWin.show()
    sys.exit(app.exec_())
   
